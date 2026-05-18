"""
Exportador DOCX — Converte o documento gerado para Microsoft Word.
Usa python-docx para criar um arquivo .docx formatado.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _is_heading(line: str) -> tuple[bool, int, str]:
    """Detecta se a linha é um título markdown e retorna (é título, nível, texto)."""
    match = re.match(r'^(#{1,4})\s+(.+)$', line.strip())
    if match:
        level = len(match.group(1))
        text = match.group(2).strip()
        return True, level, text
    return False, 0, ''


def _is_bold_line(line: str) -> tuple[bool, str]:
    """Detecta linhas inteiramente em negrito markdown **texto**."""
    match = re.match(r'^\*\*(.+)\*\*$', line.strip())
    if match:
        return True, match.group(1)
    return False, ''


def _add_paragraph_with_inline(doc: Document, text: str, style: str = 'Normal') -> None:
    """Adiciona parágrafo suportando negrito inline **texto**."""
    para = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        bold_match = re.match(r'^\*\*(.+)\*\*$', part)
        if bold_match:
            run = para.add_run(bold_match.group(1))
            run.bold = True
        else:
            para.add_run(part)


def generate_docx(title: str, content: str, norm: str = "ABNT") -> bytes:
    """
    Converte o texto markdown do documento gerado para um arquivo DOCX.

    Args:
        title: Título do projeto/documento.
        content: Texto completo gerado pelo writer_agent (pode ter markdown básico).
        norm: Norma bibliográfica usada (apenas para metadado no rodapé).

    Returns:
        Bytes do arquivo .docx gerado.
    """
    doc = Document()

    # Margens ABNT: superior 3cm, inferior 2cm, esquerda 3cm, direita 2cm
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Título do documento
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.size = Pt(14)

    doc.add_paragraph()  # linha em branco

    # Processa o conteúdo linha por linha
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Detecta cabeçalhos markdown
        is_h, level, h_text = _is_heading(stripped)
        if is_h:
            heading = doc.add_heading(h_text, level=min(level, 4))
            heading.runs[0].font.name = 'Times New Roman'
            i += 1
            continue

        # Linha de negrito completo (sem #) → heading nível 3
        is_bold, bold_text = _is_bold_line(stripped)
        if is_bold:
            heading = doc.add_heading(bold_text, level=3)
            heading.runs[0].font.name = 'Times New Roman'
            i += 1
            continue

        # Item de lista
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*[^*]+\*\*)', item_text)
            for part in parts:
                m = re.match(r'^\*\*(.+)\*\*$', part)
                if m:
                    run = para.add_run(m.group(1))
                    run.bold = True
                else:
                    para.add_run(part)
            para.runs[0].font.name = 'Times New Roman' if para.runs else None
            i += 1
            continue

        # Parágrafo normal com possível negrito inline
        _add_paragraph_with_inline(doc, stripped)
        i += 1

    # Rodapé com norma
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Documento gerado pelo AcademiaGenius · Norma {norm}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
